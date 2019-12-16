
def trans_alator(input_date):
    translation =""
    for eachletter in input_date:

        if eachletter.lower() in "aeiou":
            if eachletter.islower():
                translation = translation + "g"
            else:
                translation = translation + "G"
        else:
            translation = translation + eachletter

    return translation


'''
The comments are really good to learn

'''

# input_date = input("Enter the vowels: ")
# print(trans_alator(input_date))

''' 
 file reading using python
'''

emplyoee_details = open("employee.txt", "a")

if emplyoee_details.readable():

    for eachLine in emplyoee_details.readlines():
        print(eachLine)
else:
    print("The file is not readable")


# write to the file
    emplyoee_details.write("This is the new line\n")
    emplyoee_details.write("This is the second new line")
    emplyoee_details.close()

from docx import Document
from docx.shared import Inches

document = Document();
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('monty.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')
